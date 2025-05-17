from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
import os
from werkzeug.utils import secure_filename
import datetime
import hashlib
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import padding, rsa
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.backends import default_backend
from cryptography import x509
from cryptography.x509.oid import NameOID
import qrcode
import io
from PyPDF2 import PdfReader, PdfWriter
from docx import Document as DocxDocument
from docx.shared import Pt
import json
import rfc3161ng
import base64
import requests
from dateutil import parser
import traceback
# Importer les modules nécessaires pour la génération de PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import tempfile


app = Flask(__name__)
app.secret_key = os.urandom(24) # Nécessaire pour la gestion de session

SIGNED_DOCUMENTS_FOLDER = 'signed_documents'
if not os.path.exists(SIGNED_DOCUMENTS_FOLDER):
    os.makedirs(SIGNED_DOCUMENTS_FOLDER)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db' # Configuration de la base de données SQLite
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
migrate = Migrate(app, db)

# Clé privée et publique pour la signature (générées une fois pour la démo)
# Dans un vrai scénario, celles-ci seraient gérées de manière sécurisée et pourraient être spécifiques à l'utilisateur
private_key = rsa.generate_private_key(
    public_exponent=65537,
    key_size=2048,
    backend=default_backend()
)
public_key = private_key.public_key()

# Sérialisation des clés (optionnel, pour affichage ou stockage)
pem_private_key = private_key.private_bytes(
    encoding=serialization.Encoding.PEM,
    format=serialization.PrivateFormat.PKCS8,
    encryption_algorithm=serialization.NoEncryption()
)
pem_public_key = public_key.public_bytes(
    encoding=serialization.Encoding.PEM,
    format=serialization.PublicFormat.SubjectPublicKeyInfo
)

# Modèles de base de données
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False) # Stocker les hashs des mots de passe
    name = db.Column(db.String(100))

    def __repr__(self):
        return f'<User {self.username}>'

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    original_hash = db.Column(db.String(64), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('documents', lazy=True))

class Signature(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    document = db.relationship('Document', backref=db.backref('signatures', lazy=True, uselist=False))
    signature_hash = db.Column(db.String(512), nullable=False) # Assez grand pour une signature RSA
    signed_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    algorithm = db.Column(db.String(50), default='RSA-PSS-SHA256')

class Timestamp(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    data_description = db.Column(db.String(200))
    data_hash = db.Column(db.String(64), nullable=False)
    hash_algorithm = db.Column(db.String(50), default='SHA-256')
    timestamp_value = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    tsa_token_sim = db.Column(db.String(100)) # Jeton simulé
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True) # Peut être anonyme ou lié à un user
    user = db.relationship('User', backref=db.backref('timestamps', lazy=True))

class QrCodeLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    data_content = db.Column(db.Text, nullable=False)
    generated_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    user = db.relationship('User', backref=db.backref('qrcodes', lazy=True))

class SignatoryProcess(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False) # Lien vers le document à signer
    initiator_user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    status = db.Column(db.String(50), default='pending') # e.g., 'pending', 'in_progress', 'completed', 'cancelled'
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    completed_at = db.Column(db.DateTime)
    title = db.Column(db.String(200), nullable=False)
    document = db.relationship('Document', backref=db.backref('signatory_processes', lazy=True))
    initiator = db.relationship('User', backref=db.backref('initiated_signatory_processes', lazy=True))
    workflow_type = db.Column(db.String(50), default='sequential') # 'sequential' or 'parallel'
    audit_trail = db.Column(db.Text) # JSON string for the audit trail
    signed_document_path = db.Column(db.String(255)) # Path to the signed document

class Certificate(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('certificates', lazy=True))
    public_key_pem = db.Column(db.Text, nullable=False)
    private_key_pem = db.Column(db.Text, nullable=True) # Store encrypted private key in a real app
    issued_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    expires_at = db.Column(db.DateTime) # Optional: for certificate validity
    subject = db.Column(db.String(200)) # e.g., user's name or email
    issuer = db.Column(db.String(200)) # e.g., 'Self-Signed' or 'Your CA'

class SignatoryDocument(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    original_hash = db.Column(db.String(64), nullable=False)
    upload_date = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False) # Uploader
    user = db.relationship('User', backref=db.backref('uploaded_signatory_docs', lazy=True))
    status = db.Column(db.String(50), default='Pending') # e.g., Pending, In Progress, Completed, Rejected

class SafeDocument(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    stored_path = db.Column(db.String(255), nullable=False) # Path to the stored file
    upload_date = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', backref=db.backref('safe_documents', lazy=True))
    original_hash = db.Column(db.String(64)) # Optional: store hash for integrity check

# New models for the signing workflow
class Signer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    process_id = db.Column(db.Integer, db.ForeignKey('signatory_process.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100))
    order = db.Column(db.Integer, default=0) # Order in the signing sequence
    status = db.Column(db.String(50), default='pending') # 'pending', 'signed', 'refused', 'notified'
    signature_date = db.Column(db.DateTime)
    process = db.relationship('SignatoryProcess', backref=db.backref('signers', lazy=True, order_by='Signer.order'))
    signature_image = db.Column(db.Text) # Base64 encoded signature image
    signature_hash = db.Column(db.String(512)) # Hash of the signature
    role = db.Column(db.String(100)) # Optional role of the signer
    access_token = db.Column(db.String(64)) # Unique token for accessing the signing page

class SigningStep(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    process_id = db.Column(db.Integer, db.ForeignKey('signatory_process.id'), nullable=False)
    signer_id = db.Column(db.Integer, db.ForeignKey('signer.id'), nullable=False)
    action = db.Column(db.String(50)) # 'view', 'sign', 'decline', etc.
    timestamp = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    ip_address = db.Column(db.String(50))
    user_agent = db.Column(db.String(255))
    document_hash = db.Column(db.String(64)) # Hash of the document at this step
    details = db.Column(db.Text) # Additional details as JSON
    process = db.relationship('SignatoryProcess', backref=db.backref('steps', lazy=True))
    signer = db.relationship('Signer', backref=db.backref('steps', lazy=True))



@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user_from_db = User.query.filter_by(username=username).first()

        if user_from_db and hashlib.sha256(password.encode()).hexdigest() == user_from_db.password_hash: # Temporaire: comparaison de hash simple
            session['username'] = user_from_db.username
            session['user_id'] = user_from_db.id
            session['name'] = user_from_db.name if user_from_db.name else user_from_db.username
            return redirect(url_for('dashboard'))
        else:
            if not user_from_db:
                hashed_password = hashlib.sha256(password.encode()).hexdigest()
                new_user = User(username=username, password_hash=hashed_password, name=username.capitalize())
                db.session.add(new_user)
                db.session.commit()
                session['username'] = new_user.username
                session['user_id'] = new_user.id
                session['name'] = new_user.name
                flash(f'Compte {username} créé et connecté.', 'success')
                return redirect(url_for('dashboard'))
            return render_template('login.html', error="Nom d'utilisateur ou mot de passe incorrect.")
    return render_template('login.html')

@app.route('/dashboard')
def dashboard():
    if 'username' in session:
        return render_template('dashboard.html', name=session['name'])
    return redirect(url_for('login'))

@app.route('/logout')
def logout():
    session.pop('username', None)
    session.pop('name', None)
    session.pop('user_id', None)
    return redirect(url_for('index'))

@app.route('/services')
def services():
    
    available_services = [
        {"id": "signature", "name": "Signature Électronique", "description": "Signez numériquement vos documents en toute sécurité.", "details_url": url_for('service_detail', service_id='signature')},
        {"id": "timestamp", "name": "Horodatage Qualifié", "description": "Prouvez l'existence d'un document à une date certaine.", "details_url": url_for('service_detail', service_id='timestamp')},
        {"id": "certificates", "name": "Génération et Gestion de Certificats", "description": "Générez et gérez vos certificats numériques.", "details_url": url_for('service_detail', service_id='certificates')},
        {"id": "qrcode", "name": "Génération de Code QR Sécurisé", "description": "Générez des codes QR pour diverses applications.", "details_url": url_for('service_detail', service_id='qrcode')},
        {"id": "verification", "name": "Vérification de Signature Électronique", "description": "Vérifiez l'authenticité d'une signature numérique.", "details_url": url_for('service_detail', service_id='verification')},
        {"id": "password", "name": "Générateur de Mots de Passe Sécurisés", "description": "Créez des mots de passe forts et sécurisés pour vos comptes.", "details_url": url_for('service_detail', service_id='password')},
        {"id": "signatory", "name": "Parapheur Électronique", "description": "Gérez le flux de signature de vos documents.", "details_url": url_for('service_detail', service_id='signatory')},
        {"id": "safe", "name": "Coffre Fort Numérique", "description": "Stockez vos documents importants en toute sécurité.", "details_url": url_for('service_detail', service_id='safe')}
    ]
    return render_template('services.html', services=available_services)

@app.route('/service/<service_id>')
def service_detail(service_id):
    
    service_map = {
        "signature": {"name": "Signature Électronique", "long_description": "Notre service de signature électronique avancée permet de signer des documents PDF avec une validité légale. Il utilise des certificats qualifiés pour garantir l'authenticité et l'intégrité des signatures."},
        "timestamp": {"name": "Horodatage Qualifié", "long_description": "L'horodatage qualifié fournit une preuve irréfutable de l'existence d'un fichier ou d'une donnée à un instant T. Ce service est conforme aux réglementations eIDAS."},
        "certificates": {"name": "Génération et Gestion de Certificats", "long_description": "Générez, gérez et révoquez vos certificats numériques pour sécuriser vos communications et transactions."}, # Updated description
        "qrcode": {"name": "Génération de Code QR Sécurisé", "long_description": "Notre service de génération de codes QR vous permet de créer facilement des codes QR pour des URL, du texte, des informations de contact, et plus encore. Les codes sont générés localement pour assurer la confidentialité de vos données."},
        "verification": {"name": "Vérification de Signature Électronique", "long_description": "Ce service vous permet de téléverser un document, sa signature (en format hexadécimal) et la clé publique (en format PEM) de l'émetteur pour vérifier si la signature est authentique et si le document n'a pas été altéré depuis sa signature."},
        "password": {"name": "Générateur de Mots de Passe Sécurisés", "long_description": "Notre service de génération de mots de passe vous permet de créer des mots de passe forts et sécurisés pour vos comptes en ligne. Vous pouvez personnaliser la longueur et les caractères utilisés pour répondre à vos besoins spécifiques."},
        "signatory": {"name": "Parapheur Électronique", "long_description": "Optimisez et sécurisez le processus de validation et de signature de vos documents au sein de votre organisation."}, # New service
        "safe": {"name": "Coffre Fort Numérique", "long_description": "Stockez, organisez et accédez à vos documents importants dans un espace numérique hautement sécurisé."}
    }
    service_info = service_map.get(service_id)
    if service_info:
        
        interactive_url = None
        if service_id == 'signature':
            interactive_url = url_for('signature_service_page')
        elif service_id == 'timestamp':
            interactive_url = url_for('timestamp_service_page')
        elif service_id == 'certificates': # Updated service ID
            interactive_url = url_for('certificates_service_page') # New route
        elif service_id == 'qrcode':
            interactive_url = url_for('qr_code_service_page')
        elif service_id == 'verification':
            interactive_url = url_for('verification_service_page')
        elif service_id == 'password':
            interactive_url = url_for('password_service_page')
        elif service_id == 'signatory': # New service ID
            interactive_url = url_for('signatory_service_page') # New route
        elif service_id == 'safe': # New service ID
            interactive_url = url_for('safe_service_page') # New route
        return render_template('service_detail.html', service=service_info, interactive_url=interactive_url)
    return "Service non trouvé", 404

@app.route('/service/signature/interact', methods=['GET'])
def signature_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))
    
    return render_template('signature_service.html')

@app.route('/service/signature/action', methods=['POST'])
def signature_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    if 'document' not in request.files:
        flash('Aucun fichier sélectionné.', 'error')
        return redirect(url_for('signature_service_page'))
        
    file = request.files['document']
    if file.filename == '':
        flash('Aucun fichier sélectionné.', 'error')
        return redirect(url_for('signature_service_page'))
        
    if file and 'user_id' in session:
        file_bytes = file.read()
        original_hash = hashlib.sha256(file_bytes).hexdigest()
        
        # Enregistrer le document
        doc = Document(filename=secure_filename(file.filename), 
                      original_hash=original_hash, 
                      user_id=session['user_id'])
        db.session.add(doc)
        db.session.commit()

        # Générer la signature réelle
        signature_bytes = private_key.sign(
            file_bytes, # Les données à signer
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH
            ),
            hashes.SHA256()
        )
        signature_hex = signature_bytes.hex()

        # Enregistrer la signature
        sig = Signature(document_id=doc.id, signature_hash=signature_hex)
        db.session.add(sig)
        db.session.commit()

        message = f"Le document '{secure_filename(file.filename)}' a été signé avec succès."
        signature_details = {
            "filename": secure_filename(file.filename),
            "original_hash": original_hash,
            "signature_hex": signature_hex,
            "algorithm": sig.algorithm,
            "signed_at": sig.signed_at.strftime("%Y-%m-%d %H:%M:%S UTC"),
            "public_key_pem": pem_public_key.decode(),
            "document_id": doc.id
        }
        # Process and save the document with signature info
        signed_filename = f"signed_{secure_filename(file.filename)}"
        signed_document_path = os.path.join(SIGNED_DOCUMENTS_FOLDER, signed_filename)

        signature_info_text = f"Document: {signature_details['filename']}\n"
        signature_info_text += f"Original Hash (SHA-256): {signature_details['original_hash']}\n"
        signature_info_text += f"Signature (Hex): {signature_details['signature_hex'][:64]}...\n" # Truncate for display if too long
        signature_info_text += f"Algorithm: {signature_details['algorithm']}\n"
        signature_info_text += f"Signed At: {signature_details['signed_at']}\n"
        signature_info_text += f"Signer Public Key (PEM):\n{signature_details['public_key_pem']}"

        file_extension = os.path.splitext(file.filename)[1].lower()

        if file_extension == '.pdf':
            try:
                reader = PdfReader(io.BytesIO(file_bytes))
                writer = PdfWriter()
                for page in reader.pages:
                    writer.add_page(page)
                writer.add_metadata({
                    '/SignatureInfo': signature_info_text
                })
                with open(signed_document_path, 'wb') as f_out:
                    writer.write(f_out)
            except Exception as e:
                flash(f'Erreur lors du traitement du PDF: {e}', 'error')
                return redirect(url_for('signature_service_page'))
        elif file_extension == '.docx':
            try:
                doc_obj = DocxDocument(io.BytesIO(file_bytes))
                doc_obj.add_paragraph('\n--- Informations de Signature ---')
                for line in signature_info_text.split('\n'):
                    p = doc_obj.add_paragraph()
                    run = p.add_run(line)
                    run.font.size = Pt(10)
                doc_obj.add_paragraph('--- Fin des Informations de Signature ---')
                doc_obj.save(signed_document_path)
            except Exception as e:
                flash(f'Erreur lors du traitement du DOCX: {e}', 'error')
                return redirect(url_for('signature_service_page'))
        else:
            flash('Type de fichier non supporté pour l\'intégration de la signature. Seuls PDF et DOCX sont gérés.', 'warning')
            # Fallback: still provide signature details but no download of modified file
            return render_template('signature_service.html', message=message, signature_details=signature_details)

        download_url = url_for('download_signed_document', filename=signed_filename)
        flash(message, 'success')
        return render_template('signature_service.html', message=message, signature_details=signature_details, download_url=download_url, signed_filename=signed_filename)
    return redirect(url_for('signature_service_page'))

@app.route('/service/signatory/interact', methods=['GET'])
def signatory_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    # Fetch documents for the current user
    processes = SignatoryProcess.query.filter_by(initiator_user_id=session['user_id']).all()
    return render_template('signatory_service.html', processes=processes)

@app.route('/service/safe/interact', methods=['GET'])
def safe_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    # Fetch documents for the current user
    user_documents = SafeDocument.query.filter_by(user_id=session['user_id']).all()
    return render_template('safe_service.html', documents=user_documents)

@app.route('/download_safe_document/<int:document_id>')
def download_safe_document(document_id):
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger des documents.", "error")
        return redirect(url_for('login'))

    safe_doc = SafeDocument.query.get(document_id)

    if not safe_doc or safe_doc.user_id != session['user_id']:
        flash("Document non trouvé ou accès non autorisé.", "error")
        return redirect(url_for('safe_service_page'))

    # Get the full path from the stored_path field
    full_path = safe_doc.stored_path

    if not os.path.exists(full_path) or not os.path.isfile(full_path):
         flash("Fichier physique non trouvé.", "error")
         return redirect(url_for('safe_service_page'))

    # Additional check to prevent path traversal
    safe_folder = 'safe_documents'
    if not os.path.commonprefix((os.path.realpath(full_path), os.path.realpath(safe_folder))) == os.path.realpath(safe_folder):
         flash("Accès non autorisé au fichier.", "error")
         return redirect(url_for('safe_service_page'))

    return send_file(full_path, as_attachment=True, download_name=safe_doc.filename)

@app.route('/service/timestamp/interact', methods=['GET'])
def timestamp_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    return render_template('timestamp_service.html')

@app.route('/service/timestamp/action', methods=['POST'])
def timestamp_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))
    
    data_text = request.form.get('data_text')
    data_file = request.files.get('data_file')
    
    if not data_text and (not data_file or data_file.filename == ''):
        flash('Veuillez fournir du texte ou un fichier à horodater.', 'error')
        return render_template('timestamp_service.html', error='Veuillez fournir du texte ou un fichier à horodater.')

    timestamp_str = datetime.datetime.now().isoformat() + "Z"
    data_hash = None
    data_description = ""
    filename = None
    original_data = None

    if data_file and data_file.filename != '':
        filename = secure_filename(data_file.filename)
        file_bytes = data_file.read()
        original_data = file_bytes
        data_hash = hashlib.sha256(file_bytes).hexdigest()
        data_file.seek(0) # Important si le fichier doit être relu ou sauvegardé
        data_description = f"Fichier : {filename}"
    elif data_text:
        original_data = data_text.encode('utf-8')
        data_hash = hashlib.sha256(original_data).hexdigest()
        data_description = f"Texte : {data_text[:50]}{'...' if len(data_text) > 50 else ''}"
    
    if data_hash:
        # Toujours utiliser la simulation d'horodatage
        current_time = datetime.datetime.utcnow()
        tsa_timestamp = current_time
        
        # Créer une signature basée sur le hash et l'horodatage
        timestamp_data = f"{data_hash}:{current_time.isoformat()}"
        signature = private_key.sign(
            timestamp_data.encode(),
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH
            ),
            hashes.SHA256()
        )
        
        # Construire un jeton d'horodatage structuré
        timestamp_token = {
            "version": "1.0",
            "hash_algorithm": "SHA-256",
            "hash_value": data_hash,
            "timestamp": current_time.isoformat(),
            "status": "granted",
            "tsa_name": "TrustSecure Demo TSA",
            "tsa_signature": base64.b64encode(signature).decode('utf-8')
        }
        
        # Encoder le jeton en JSON puis base64 pour le stockage
        token_json = json.dumps(timestamp_token)
        tsa_token_b64 = base64.b64encode(token_json.encode('utf-8')).decode('utf-8')
        
        # Générer aussi un jeton dans l'ancien format pour la compatibilité
        tsa_token_sim = f"TSA_TOKEN_SIM_{hashlib.md5((data_hash + current_time.isoformat()).encode()).hexdigest().upper()}"
        
        
        is_rfc3161_success = False
        tsa_note = """Note: Ce service utilise un mode simulation d'horodatage qui est cryptographiquement sécurisé mais qui n'est pas connecté à une Autorité d'Horodatage externe.
        
        Le jeton généré est néanmoins cryptographiquement vérifiable avec la clé publique fournie et peut être utilisé pour prouver l'existence de la donnée au moment indiqué."""
        
        # Enregistrer l'horodatage
        ts = Timestamp(
            data_description=data_description,
            data_hash=data_hash,
            hash_algorithm='SHA-256',
            timestamp_value=tsa_timestamp,
            tsa_token_sim=tsa_token_b64,  # On utilise le jeton formaté en base64
            user_id=session.get('user_id')
        )
        db.session.add(ts)
        db.session.commit()

        message = f"La donnée a été horodatée avec succès."
        timestamp_details = {
            "id": ts.id,
            "description": data_description,
            "hash_algorithm": ts.hash_algorithm,
            "hash_value": ts.data_hash,
            "timestamp": ts.timestamp_value.strftime("%Y-%m-%d %H:%M:%S.%f")[:-3] + "Z",
            "tsa_token": ts.tsa_token_sim,
            "rfc3161_success": is_rfc3161_success,
            "tsa_note": tsa_note
        }
        
        if filename:
            timestamp_details["filename"] = filename

        flash(message, 'success')
        return render_template('timestamp_service.html', message=message, timestamp_details=timestamp_details)

@app.route('/service/identity/interact', methods=['GET'])
def identity_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    return render_template('identity_service.html')

@app.route('/service/identity/action', methods=['POST'])
def identity_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    full_name = request.form.get('fullName')
    id_number = request.form.get('idNumber')
    id_document = request.files.get('idDocument')

    if not full_name or not id_number:
        flash('Veuillez remplir tous les champs obligatoires.', 'error')
        return render_template('identity_service.html', error='Veuillez remplir tous les champs obligatoires.', fullName=full_name, idNumber=id_number)

    # Simulation de la vérification
    verification_status = "En attente de validation manuelle"
    transaction_id = f"IDV_SIM_{hashlib.md5((full_name + id_number + datetime.datetime.now().isoformat()).encode()).hexdigest()[:12].upper()}"
    document_received_message = ""

    # Logique de simulation simple
    if len(full_name.split()) >= 2 and len(id_number) > 5:
        if id_document and id_document.filename != '':
            
            if any(char.isdigit() for char in full_name):
                verification_status = "Échouée (Nom invalide)"
            elif not id_number.isalnum():
                verification_status = "Échouée (Numéro d'ID invalide)"
            else:
                verification_status = "Vérifiée (Simulation)"
            document_received_message = f"Le document '{secure_filename(id_document.filename)}' a été reçu et pris en compte."
        else:
            # Moins d'informations, donc statut différent
            verification_status = "Partiellement vérifiée (Document manquant - Simulation)"
    else:
        verification_status = "Échouée (Données fournies insuffisantes)"

    message = f"Le traitement de la vérification d'identité pour '{full_name}' est terminé."
    identity_details = {
        "fullName": full_name,
        "idNumber": id_number,
        "document_message": document_received_message,
        "status": verification_status,
        "transaction_id": transaction_id,
        "verification_date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC")
    }
    
    flash(message, 'message')
    return render_template('identity_service.html', message=message, identity_details=identity_details, fullName=full_name, idNumber=id_number)

@app.route('/service/qrcode/interact', methods=['GET'])
def qr_code_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "warning")
        
    return render_template('qr_code_service.html')

@app.route('/service/qrcode/generate', methods=['POST'])
def qr_code_generate_action():
    data_to_encode = request.form.get('qr_data')
    if not data_to_encode:
        flash('Veuillez entrer des données à encoder dans le QR code.', 'error')
        return render_template('qr_code_service.html', error='Veuillez entrer des données à encoder.')

    # Génération du QR Code
    img = qrcode.make(data_to_encode)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)

    # Enregistrer le log de génération
    qr_log = QrCodeLog(
        data_content=data_to_encode,
        user_id=session.get('user_id')
    )
    db.session.add(qr_log)
    db.session.commit()

    flash('Code QR généré avec succès!', 'success')
   
    session['last_qr_id'] = qr_log.id
    return redirect(url_for('qr_code_service_page_with_image', qr_id=qr_log.id))

@app.route('/service/qrcode/image/<int:qr_id>')
def serve_qr_code_image(qr_id):
    qr_log_entry = QrCodeLog.query.get_or_404(qr_id)

        
    img = qrcode.make(qr_log_entry.data_content)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return send_file(buf, mimetype='image/png', as_attachment=False, download_name=f'qrcode_{qr_id}.png')

@app.route('/service/qrcode/interact/image/<int:qr_id>')
def qr_code_service_page_with_image(qr_id):
    qr_log_entry = QrCodeLog.query.get_or_404(qr_id)
    return render_template('qr_code_service.html', qr_image_url=url_for('serve_qr_code_image', qr_id=qr_id), data_encoded=qr_log_entry.data_content)

@app.route('/service/verification/interact', methods=['GET'])
def verification_service_page():
    
    return render_template('verification_service.html')

@app.route('/service/verification/action', methods=['POST'])
def verification_service_action():
    if 'document' not in request.files or \
       'signature_hex' not in request.form or \
       'public_key_pem' not in request.form:
        flash('Tous les champs sont requis.', 'error')
        return redirect(url_for('verification_service_page'))

    document_file = request.files['document']
    signature_hex = request.form['signature_hex'].strip()
    public_key_pem_str = request.form['public_key_pem'].strip()

    if document_file.filename == '':
        flash('Aucun document sélectionné.', 'error')
        return redirect(url_for('verification_service_page'))

    if not signature_hex:
        flash('La signature hexadécimale est requise.', 'error')
        return redirect(url_for('verification_service_page'))

    if not public_key_pem_str:
        flash('La clé publique PEM est requise.', 'error')
        return redirect(url_for('verification_service_page'))

    verification_result = {'filename': secure_filename(document_file.filename)}

    try:
        document_bytes = document_file.read()
        document_hash = hashlib.sha256(document_bytes).hexdigest()
        verification_result['document_hash'] = document_hash

        # Charger la clé publique
        public_key_to_verify = serialization.load_pem_public_key(
            public_key_pem_str.encode(),
            backend=default_backend()
        )

        # Décoder la signature hexadécimale en bytes
        try:
            signature_bytes = bytes.fromhex(signature_hex)
        except ValueError:
            flash('Format de signature hexadécimale invalide.', 'error')
            verification_result['valid'] = False
            verification_result['error_message'] = 'Format de signature hexadécimale invalide.'
            return render_template('verification_service.html', verification_result=verification_result)

        # Vérifier la signature
        public_key_to_verify.verify(
            signature_bytes,
            document_bytes, 
            padding.PSS(
                mgf=padding.MGF1(hashes.SHA256()),
                salt_length=padding.PSS.MAX_LENGTH
            ),
            hashes.SHA256() 
        )
        verification_result['valid'] = True
        flash('La signature est valide.', 'success')

    except ValueError as e:
        # Erreur de format PEM ou autre problème de clé
        verification_result['valid'] = False
        verification_result['error_message'] = f"Erreur avec la clé publique : {str(e)}"
        flash(f"Erreur avec la clé publique : {str(e)}", 'danger')
    except Exception as e:
        # Cela inclut InvalidSignature de cryptography
        verification_result['valid'] = False
        verification_result['error_message'] = f"La signature est invalide ou une erreur est survenue : {str(e)}"
        flash(f"La signature est invalide ou une erreur est survenue : {str(e)}", 'danger')

    return render_template('verification_service.html', verification_result=verification_result)

@app.route('/service/password/interact', methods=['GET'])
def password_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    return render_template('password_service.html')

@app.route('/service/password/action', methods=['POST'])
def password_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    # Placeholder for password generation logic
    generated_password = "GeneratedPassword123!" # Simulate password generation
    flash("Mot de passe généré (simulation).", "success")
    return render_template('password_service.html', generated_password=generated_password)

@app.route('/service/certificates/interact', methods=['GET'])
def certificates_service_page():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    user_certificates = Certificate.query.filter_by(user_id=session['user_id']).all()
    return render_template('certificates_service.html', certificates=user_certificates)

@app.route('/service/certificates/action', methods=['POST'])
def certificates_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    subject_name = request.form.get('subject_name')
    if not subject_name:
        flash("Le nom du sujet est requis.", "error")
        return redirect(url_for('certificates_service_page'))

    try:
        # Générer une nouvelle clé privée pour ce certificat
        new_private_key = rsa.generate_private_key(
            public_exponent=65537,
            key_size=2048,
            backend=default_backend()
        )

        # Créer un certificat auto-signé
        subject = issuer = x509.Name([
            x509.NameAttribute(NameOID.COMMON_NAME, subject_name),
        ])
        cert = x509.CertificateBuilder().subject_name(subject)
        cert = cert.issuer_name(issuer)
        cert = cert.public_key(new_private_key.public_key())
        cert = cert.serial_number(x509.random_serial_number())
        cert = cert.not_valid_before(datetime.datetime.utcnow())
        cert = cert.not_valid_after(datetime.datetime.utcnow() + datetime.timedelta(days=365)) # Valide pour 1 an
        cert = cert.add_extension(
            x509.SubjectAlternativeName([x509.DNSName("localhost")]),
            critical=False,
        )
        cert = cert.sign(new_private_key, hashes.SHA256(), default_backend())

        # Sérialiser la clé privée et le certificat
        pem_private = new_private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        )
        pem_cert = cert.public_bytes(serialization.Encoding.PEM)

        # Enregistrer dans la base de données
        new_certificate = Certificate(
            user_id=session['user_id'],
            public_key_pem=new_private_key.public_key().public_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PublicFormat.SubjectPublicKeyInfo
            ).decode(),
            private_key_pem=pem_private.decode(),
            issued_at=datetime.datetime.utcnow(),
            expires_at=datetime.datetime.utcnow() + datetime.timedelta(days=365),
            subject=subject_name,
            issuer=subject_name  # Auto-signé
        )
        db.session.add(new_certificate)
        db.session.commit()

        flash(f'Certificat pour "{subject_name}" généré avec succès.', 'success')

    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la génération du certificat : {e}', 'error')

    return redirect(url_for('certificates_service_page'))

@app.route('/service/signatory/action', methods=['POST'])
def signatory_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    if 'document' not in request.files:
        flash('Aucun fichier sélectionné pour le processus de signature.', 'error')
        return redirect(url_for('signatory_service_page'))

    file = request.files['document']
    process_title = request.form.get('process_title')

    if file.filename == '' or not process_title:
        flash('Veuillez sélectionner un fichier et fournir un titre pour le processus.', 'error')
        return redirect(url_for('signatory_service_page'))

    if file and 'user_id' in session:
        try:
            file_bytes = file.read()
            original_hash = hashlib.sha256(file_bytes).hexdigest()
            
            # Create a directory to store the original and signed documents
            doc_directory = os.path.join('signatory_documents', str(session['user_id']))
            if not os.path.exists(doc_directory):
                os.makedirs(doc_directory)
                
            # Save the original document
            secure_filename_val = secure_filename(file.filename)
            original_file_path = os.path.join(doc_directory, secure_filename_val)
            with open(original_file_path, 'wb') as f:
                f.write(file_bytes)

            # Enregistrer le document dans la base de données
            doc = Document(
                filename=secure_filename_val,
                original_hash=original_hash,
                user_id=session['user_id']
            )
            db.session.add(doc)
            db.session.flush() # Pour obtenir l'ID du document avant le commit

            # Créer le processus de signature
            signatory_process = SignatoryProcess(
                document_id=doc.id,
                initiator_user_id=session['user_id'],
                title=process_title,
                status='pending', # Initial status
                workflow_type='sequential'
            )
            db.session.add(signatory_process)
            db.session.flush()
            
            # Try to extract potential signers from the document
            extracted_signers = []
            file_extension = os.path.splitext(file.filename)[1].lower()
            
            try:
                if file_extension == '.pdf':
                    # Attempt to extract text from PDF
                    file.seek(0)
                    pdf = PdfReader(io.BytesIO(file_bytes))
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text()
                    
                    # Simple pattern matching for potential signers (very basic)
                    # Look for patterns like "Signature:" or "Sign:" followed by names
                    import re
                    signature_patterns = [
                        r'Signature[s]?[\s:]+([\w\s]+)',
                        r'Sign[é]?[\s:]+([\w\s]+)',
                        r'Approbation[\s:]+([\w\s]+)',
                        r'Approuv[é]?[\s:]+([\w\s]+)'
                    ]
                    
                    for pattern in signature_patterns:
                        matches = re.findall(pattern, text)
                        for match in matches:
                            name = match.strip()
                            if len(name) > 3 and name not in [s['name'] for s in extracted_signers]:
                                extracted_signers.append({'name': name, 'role': 'Signataire'})
                
                elif file_extension == '.docx':
                    # Attempt to extract text from DOCX
                    file.seek(0)
                    doc_obj = DocxDocument(io.BytesIO(file_bytes))
                    text = ""
                    for para in doc_obj.paragraphs:
                        text += para.text + "\n"
                    
                    # Simple pattern matching for potential signers
                    import re
                    signature_patterns = [
                        r'Signature[s]?[\s:]+([\w\s]+)',
                        r'Sign[é]?[\s:]+([\w\s]+)',
                        r'Approbation[\s:]+([\w\s]+)',
                        r'Approuv[é]?[\s:]+([\w\s]+)'
                    ]
                    
                    for pattern in signature_patterns:
                        matches = re.findall(pattern, text)
                        for match in matches:
                            name = match.strip()
                            if len(name) > 3 and name not in [s['name'] for s in extracted_signers]:
                                extracted_signers.append({'name': name, 'role': 'Signataire'})
            except Exception as e:
                # If extraction fails, we'll use default signers ####FIRAS
                print(f"Error extracting signers: {e}")
            
            # If no signers found, create a default signer ##########FIRAS
            if not extracted_signers:
                extracted_signers = [
                    {'name': 'Signataire 1', 'role': 'Signataire Principal'},
                    {'name': 'Signataire 2', 'role': 'Validateur'}
                ]
            
            # Create signers
            for i, signer_info in enumerate(extracted_signers):
                signer = Signer(
                    process_id=signatory_process.id,
                    name=signer_info['name'],
                    role=signer_info['role'],
                    order=i,
                    status='pending',
                    access_token=hashlib.sha256(f"{signatory_process.id}_{i}_{datetime.datetime.now().isoformat()}".encode()).hexdigest()
                )
                db.session.add(signer)
            
            # Initialize audit trail ##########FIRAS
            audit_data = {
                'process_id': signatory_process.id,
                'title': process_title,
                'created_at': datetime.datetime.now().isoformat(),
                'initiator': session['username'],
                'document': {
                    'filename': secure_filename_val,
                    'hash': original_hash
                },
                'events': [
                    {
                        'action': 'process_created',
                        'timestamp': datetime.datetime.now().isoformat(),
                        'user': session['username'],
                        'details': 'Process created'
                    }
                ]
            }
            signatory_process.audit_trail = json.dumps(audit_data)
            
            db.session.commit()

            flash(f'Processus de signature "{process_title}" initié avec succès pour le document "{secure_filename_val}".', 'success')
            
            # Redirect to signing process details page ##########FIRAS
            return redirect(url_for('view_signatory_process', process_id=signatory_process.id))

        except Exception as e:
            db.session.rollback()
            flash(f"Erreur lors de l'initiation du processus de signature : {e}", 'error')

    return redirect(url_for('signatory_service_page'))

@app.route('/service/safe/action', methods=['POST'])
def safe_service_action():
    if 'username' not in session:
        flash("Veuillez vous connecter pour utiliser ce service.", "error")
        return redirect(url_for('login'))

    if 'safe_document' not in request.files:
        flash('Aucun fichier sélectionné pour le coffre-fort numérique.', 'error')
        return redirect(url_for('safe_service_page'))

    file = request.files['safe_document']

    if file.filename == '':
        flash('Aucun fichier sélectionné.', 'error')
        return redirect(url_for('safe_service_page'))

    if file and 'user_id' in session:
        try:
            filename = secure_filename(file.filename)
            
            SAFE_DOCUMENTS_FOLDER = 'safe_documents'
            if not os.path.exists(SAFE_DOCUMENTS_FOLDER):
                os.makedirs(SAFE_DOCUMENTS_FOLDER)

            stored_path = os.path.join(SAFE_DOCUMENTS_FOLDER, filename)
            file.save(stored_path)

            # Calculate hash of the original file before potential encryption
            file.seek(0) 
            original_hash = hashlib.sha256(file.read()).hexdigest()

            # Enregistrer dans la base de données
            new_safe_document = SafeDocument(
                user_id=session['user_id'],
                filename=filename,
                stored_path=stored_path,
                original_hash=original_hash
            )
            db.session.add(new_safe_document)
            db.session.commit()

            flash(f'Document "{filename}" stocké dans le coffre-fort numérique.', 'success')

        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors du stockage du document : {e}', 'error')
    
    return redirect(url_for('safe_service_page'))

@app.route('/download_signed/<filename>')
def download_signed_document(filename):
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger des documents.", "error")
        return redirect(url_for('login'))
    
    # Sécurisation du chemin
    safe_path = os.path.join(SIGNED_DOCUMENTS_FOLDER, filename)
    if not os.path.exists(safe_path) or not os.path.isfile(safe_path):
        flash("Fichier non trouvé.", "error")
        return redirect(url_for('services')) # ou une autre page d'erreur appropriée
    
    # Vérifier que le fichier est bien dans le dossier attendu (prévention du path traversal)
    if os.path.commonprefix((os.path.realpath(safe_path), os.path.realpath(SIGNED_DOCUMENTS_FOLDER))) != os.path.realpath(SIGNED_DOCUMENTS_FOLDER):
        flash("Accès non autorisé au fichier.", "error")
        return redirect(url_for('services'))

    return send_file(safe_path, as_attachment=True)

@app.route('/download_signature/<int:document_id>')
def download_signature_hex(document_id):
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger des signatures.", "error")
        return redirect(url_for('login'))

    signature = Signature.query.filter_by(document_id=document_id).first()
    if not signature:
        flash("Signature non trouvée.", "error")
        return redirect(url_for('services'))

    # Créer un fichier temporaire ou un objet BytesIO pour la signature hex
    signature_content = signature.signature_hash.encode('utf-8')
    return send_file(io.BytesIO(signature_content), as_attachment=True, download_name=f'signature_{document_id}.hex', mimetype='text/plain')

@app.route('/download_certificate/<int:cert_id>')
def download_certificate(cert_id):
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger des certificats.", "error")
        return redirect(url_for('login'))

    certificate = Certificate.query.get_or_404(cert_id)
    
    # Verify the user owns this certificate
    if certificate.user_id != session['user_id']:
        flash("Accès non autorisé à ce certificat.", "error")
        return redirect(url_for('certificates_service_page'))
    
    # Create certificate content
    cert_content = f"Subject: {certificate.subject}\n"
    cert_content += f"Issuer: {certificate.issuer}\n"
    cert_content += f"Issued At: {certificate.issued_at.strftime('%Y-%m-%d %H:%M:%S')}\n"
    if certificate.expires_at:
        cert_content += f"Expires At: {certificate.expires_at.strftime('%Y-%m-%d %H:%M:%S')}\n"
    cert_content += f"\nPublic Key:\n{certificate.public_key_pem}\n"
    
    # Optionally include private key (in a real app, this would be heavily secured)
    if certificate.private_key_pem:
        cert_content += f"\nPrivate Key (SENSITIVE):\n{certificate.private_key_pem}\n"
    
    # Serve the certificate as a downloadable file
    return send_file(
        io.BytesIO(cert_content.encode('utf-8')),
        as_attachment=True,
        download_name=f'certificate_{certificate.subject}_{cert_id}.pem',
        mimetype='text/plain'
    )

@app.route('/service/signatory/process/<int:process_id>')
def view_signatory_process(process_id):
    """View the details of a signing process and its signers"""
    if 'username' not in session:
        flash("Veuillez vous connecter pour accéder à ce processus.", "error")
        return redirect(url_for('login'))
    
    # Get the process
    process = SignatoryProcess.query.get_or_404(process_id)
    
    # Check if the user is authorized to view this process
    if process.initiator_user_id != session['user_id']:
        flash("Vous n'êtes pas autorisé à accéder à ce processus.", "error")
        return redirect(url_for('signatory_service_page'))
    
    # Get related data
    document = process.document
    signers = process.signers
    
    # Create signing links
    signing_links = []
    for signer in signers:
        signing_url = url_for('signing_page', token=signer.access_token, _external=True)
        signing_links.append({
            'name': signer.name,
            'status': signer.status,
            'url': signing_url
        })
    
    return render_template(
        'signatory_process_detail.html', 
        process=process, 
        document=document, 
        signers=signers,
        signing_links=signing_links
    )

@app.route('/signing/<token>')
def signing_page(token):
    """Public page for a signer to view and sign a document"""
    # Find the signer by token
    signer = Signer.query.filter_by(access_token=token).first_or_404()
    process = signer.process
    document = process.document
    
    # Record view in audit trail
    signing_step = SigningStep(
        process_id=process.id,
        signer_id=signer.id,
        action='view',
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string,
        document_hash=document.original_hash
    )
    db.session.add(signing_step)
    
    # Update audit trail
    audit_data = json.loads(process.audit_trail) if process.audit_trail else {'events': []}
    audit_data['events'].append({
        'action': 'document_viewed',
        'timestamp': datetime.datetime.now().isoformat(),
        'signer': signer.name,
        'details': f"Document viewed by {signer.name}"
    })
    process.audit_trail = json.dumps(audit_data)
    db.session.commit()
    
    # Get the document path
    doc_directory = os.path.join('signatory_documents', str(process.initiator_user_id))
    document_path = os.path.join(doc_directory, document.filename)
    
    return render_template(
        'signing_page.html',
        signer=signer,
        process=process,
        document=document,
        document_path=document_path
    )

@app.route('/signing/<token>/sign', methods=['POST'])
def sign_document(token):
    """Handle document signing"""
    signer = Signer.query.filter_by(access_token=token).first_or_404()
    process = signer.process
    document = process.document
    
    # Check if the process is still active
    if process.status not in ['pending', 'in_progress']:
        flash("Ce processus de signature n'est plus actif.", "error")
        return redirect(url_for('signing_page', token=token))
    
    # Check if this signer is allowed to sign now (respecting sequential workflow)
    if process.workflow_type == 'sequential':
        # Find the current active signer (lowest order with pending status)
        current_active_signer = Signer.query.filter_by(
            process_id=process.id, 
            status='pending'
        ).order_by(Signer.order).first()
        
        if current_active_signer.id != signer.id:
            flash("Ce n'est pas encore votre tour de signer ce document.", "error")
            return redirect(url_for('signing_page', token=token))
    
    # Get signature data
    signature_image = request.form.get('signature_image', '')  # Base64 encoded image
    
    # Update signer status
    signer.status = 'signed'
    signer.signature_date = datetime.datetime.utcnow()
    signer.signature_image = signature_image
    
    # Calculate signature hash
    signature_data = f"{document.original_hash}_{signer.id}_{signer.signature_date.isoformat()}"
    signer.signature_hash = hashlib.sha256(signature_data.encode()).hexdigest()
    
    # Record signing step
    signing_step = SigningStep(
        process_id=process.id,
        signer_id=signer.id,
        action='sign',
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string,
        document_hash=document.original_hash,
        details=json.dumps({
            'signature_hash': signer.signature_hash,
            'signature_time': signer.signature_date.isoformat()
        })
    )
    db.session.add(signing_step)
    
    # Update process status
    all_signed = all(s.status == 'signed' for s in process.signers)
    if all_signed:
        process.status = 'completed'
        process.completed_at = datetime.datetime.utcnow()
        
        # Generate final signed document
        signed_file_path = generate_signed_document(process, document)
        process.signed_document_path = signed_file_path
        
        # Generate audit trail document
        generate_audit_trail_document(process)
    else:
        # If sequential and there are more signers, update process status
        if process.workflow_type == 'sequential':
            process.status = 'in_progress'
    
    # Update audit trail
    audit_data = json.loads(process.audit_trail) if process.audit_trail else {'events': []}
    audit_data['events'].append({
        'action': 'document_signed',
        'timestamp': datetime.datetime.now().isoformat(),
        'signer': signer.name,
        'details': f"Document signed by {signer.name}"
    })
    
    if all_signed:
        audit_data['events'].append({
            'action': 'process_completed',
            'timestamp': datetime.datetime.now().isoformat(),
            'details': "Document signing process completed"
        })
    
    process.audit_trail = json.dumps(audit_data)
    db.session.commit()
    
    flash("Votre signature a été enregistrée avec succès.", "success")
    return redirect(url_for('signing_confirmation', token=token))

@app.route('/signing/<token>/confirmation')
def signing_confirmation(token):
    """Show confirmation page after signing"""
    signer = Signer.query.filter_by(access_token=token).first_or_404()
    process = signer.process
    
    return render_template(
        'signing_confirmation.html',
        signer=signer,
        process=process
    )

@app.route('/service/signatory/process/<int:process_id>/download')
def download_signatory_document(process_id):
    """Download the signed document"""
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger ce document.", "error")
        return redirect(url_for('login'))
    
    process = SignatoryProcess.query.get_or_404(process_id)
    
    # Check if the user is authorized
    if process.initiator_user_id != session['user_id']:
        flash("Vous n'êtes pas autorisé à télécharger ce document.", "error")
        return redirect(url_for('signatory_service_page'))
    
    # Check if the document is signed and available
    if not process.signed_document_path or not os.path.exists(process.signed_document_path):
        flash("Le document signé n'est pas disponible.", "error")
        return redirect(url_for('view_signatory_process', process_id=process_id))
    
    return send_file(process.signed_document_path, as_attachment=True)

@app.route('/service/signatory/process/<int:process_id>/audit-trail')
def download_audit_trail(process_id):
    """Download the audit trail document"""
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger la piste d'audit.", "error")
        return redirect(url_for('login'))
    
    process = SignatoryProcess.query.get_or_404(process_id)
    
    # Check if the user is authorized
    if process.initiator_user_id != session['user_id']:
        flash("Vous n'êtes pas autorisé à télécharger la piste d'audit.", "error")
        return redirect(url_for('signatory_service_page'))
    
    # Generate an audit trail document if it doesn't exist
    audit_trail_path = os.path.join(
        'signatory_documents', 
        str(process.initiator_user_id),
        f'audit_trail_{process.id}.pdf'
    )
    
    if not os.path.exists(audit_trail_path):
        generate_audit_trail_document(process)
    
    if not os.path.exists(audit_trail_path):
        flash("La piste d'audit n'est pas disponible.", "error")
        return redirect(url_for('view_signatory_process', process_id=process_id))
    
    return send_file(audit_trail_path, as_attachment=True, download_name=f'audit_trail_{process.id}.pdf')

@app.cli.command('init-db')
def init_db_command():
    """Initialise la base de données."""
    db.create_all()
    print('Base de données initialisée.')

def generate_signed_document(process, document):
    """Generate a signed version of the document with signature information"""
    # Get the original document path
    doc_directory = os.path.join('signatory_documents', str(process.initiator_user_id))
    original_file_path = os.path.join(doc_directory, document.filename)
    
    if not os.path.exists(original_file_path):
        return None
    
    file_extension = os.path.splitext(document.filename)[1].lower()
    signed_filename = f"signed_{document.filename}"
    signed_file_path = os.path.join(doc_directory, signed_filename)
    
    try:
        if file_extension == '.pdf':
            with open(original_file_path, 'rb') as file:
                reader = PdfReader(file)
                writer = PdfWriter()
                
                # Copy all pages from the original document
                for page in reader.pages:
                    writer.add_page(page)
                
                # Add signature information
                signature_info = {
                    '/Title': f"Document signé: {document.filename}",
                    '/Author': "TrustSecure - Service de Signature Électronique",
                    '/Subject': f"Processus de signature: {process.title}",
                    '/Keywords': "signature électronique, document signé, parapheur électronique",
                    '/CreationDate': datetime.datetime.now().strftime("%Y%m%d%H%M%S"),
                    '/SignatureInfo': generate_signature_info_text(process)
                }
                writer.add_metadata(signature_info)
                
                
                
                # Save the signed document  
                with open(signed_file_path, 'wb') as output_file:
                    writer.write(output_file)
                
        elif file_extension == '.docx':
            doc = DocxDocument(original_file_path)
            
            # Add signature section at the end of the document
            doc.add_paragraph("\n")
            doc.add_heading("Informations de Signature", level=1)
            doc.add_paragraph(f"Document: {document.filename}")
            doc.add_paragraph(f"Processus: {process.title}")
            doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("\n")
            
            # Add signature blocks for each signer
            doc.add_heading("Signatures", level=2)
            for signer in process.signers:
                p = doc.add_paragraph()
                p.add_run(f"Signataire: {signer.name}").bold = True
                p.add_run(f"\nRôle: {signer.role}")
                p.add_run(f"\nStatut: {signer.status}")
                if signer.signature_date:
                    p.add_run(f"\nDate de signature: {signer.signature_date.strftime('%Y-%m-%d %H:%M:%S')}")
                p.add_run(f"\nHash de signature: {signer.signature_hash}")
                doc.add_paragraph("\n")
            
            # Save the signed document
            doc.save(signed_file_path)
        
        return signed_file_path
    except Exception as e:
        print(f"Error generating signed document: {e}")
        return None

def generate_signature_info_text(process):
    """Generate a text representation of the signature information"""
    info_text = f"Document: {process.title}\n"
    info_text += f"Processus initié par: {process.initiator.username}\n"
    info_text += f"Date d'initiation: {process.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n"
    if process.completed_at:
        info_text += f"Date de complétion: {process.completed_at.strftime('%Y-%m-%d %H:%M:%S')}\n"
    info_text += f"Statut: {process.status}\n\n"
    
    info_text += "Signataires:\n"
    for signer in process.signers:
        info_text += f"- {signer.name} ({signer.role}): {signer.status}"
        if signer.signature_date:
            info_text += f" le {signer.signature_date.strftime('%Y-%m-%d %H:%M:%S')}"
        info_text += f"\n  Hash de signature: {signer.signature_hash}\n"
    
    return info_text

def generate_audit_trail_document(process):
    """Generate an audit trail document (PDF)"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    except ImportError:
        # If reportlab is not installed, use a simple text file
        doc_directory = os.path.join('signatory_documents', str(process.initiator_user_id))
        audit_trail_path = os.path.join(doc_directory, f'audit_trail_{process.id}.txt')
        
        audit_data = json.loads(process.audit_trail) if process.audit_trail else {}
        with open(audit_trail_path, 'w') as f:
            f.write(f"Audit Trail for Process: {process.title}\n")
            f.write(f"Document: {process.document.filename}\n")
            f.write(f"Process ID: {process.id}\n")
            f.write(f"Initiator: {process.initiator.username}\n")
            f.write(f"Created: {process.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Status: {process.status}\n\n")
            
            f.write("Signers:\n")
            for signer in process.signers:
                f.write(f"- {signer.name} ({signer.role}): {signer.status}")
                if signer.signature_date:
                    f.write(f" on {signer.signature_date.strftime('%Y-%m-%d %H:%M:%S')}")
                f.write(f"\n  Signature Hash: {signer.signature_hash}\n")
            
            f.write("\nEvents:\n")
            for event in audit_data.get('events', []):
                f.write(f"- {event.get('timestamp')}: {event.get('action')} - {event.get('details')}\n")
            
            f.write("\nDocument Hash: {process.document.original_hash}\n")
        
        return audit_trail_path
    
    # Using reportlab for a nicer PDF
    doc_directory = os.path.join('signatory_documents', str(process.initiator_user_id))
    if not os.path.exists(doc_directory):
        os.makedirs(doc_directory)
    
    audit_trail_path = os.path.join(doc_directory, f'audit_trail_{process.id}.pdf')
    
    doc = SimpleDocTemplate(audit_trail_path, pagesize=letter)
    elements = []
    
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    heading2_style = styles['Heading2']
    normal_style = styles['Normal']
    
    # Add title
    elements.append(Paragraph(f"Audit Trail - {process.title}", title_style))
    elements.append(Spacer(1, 12))
    
    # Add document info
    elements.append(Paragraph("Document Information", heading2_style))
    elements.append(Spacer(1, 6))
    
    doc_info = [
        ["Document", process.document.filename],
        ["Process ID", str(process.id)],
        ["Initiator", process.initiator.username],
        ["Created", process.created_at.strftime('%Y-%m-%d %H:%M:%S')],
        ["Status", process.status]
    ]
    
    t = Table(doc_info, colWidths=[100, 400])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))
    
    # Add signers info
    elements.append(Paragraph("Signers", heading2_style))
    elements.append(Spacer(1, 6))
    
    signer_data = [["Name", "Role", "Status", "Signature Date", "Signature Hash"]]
    for signer in process.signers:
        sig_date = signer.signature_date.strftime('%Y-%m-%d %H:%M:%S') if signer.signature_date else "N/A"
        sig_hash = signer.signature_hash[:12] + "..." if signer.signature_hash else "N/A"
        signer_data.append([signer.name, signer.role, signer.status, sig_date, sig_hash])
    
    t = Table(signer_data, colWidths=[80, 80, 60, 100, 120])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))
    
    # Add events
    elements.append(Paragraph("Events", heading2_style))
    elements.append(Spacer(1, 6))
    
    audit_data = json.loads(process.audit_trail) if process.audit_trail else {}
    event_data = [["Timestamp", "Action", "Details"]]
    for event in audit_data.get('events', []):
        event_data.append([
            event.get('timestamp', 'N/A'),
            event.get('action', 'N/A'),
            event.get('details', 'N/A')
        ])
    
    t = Table(event_data, colWidths=[150, 100, 250])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))
    
    # Add hash info
    elements.append(Paragraph("Document Integrity", heading2_style))
    elements.append(Spacer(1, 6))
    elements.append(Paragraph(f"Original Document Hash (SHA-256): {process.document.original_hash}", normal_style))
    
    # Build the PDF
    doc.build(elements)
    
    return audit_trail_path

@app.route('/service/timestamp/generate_pdf/<int:timestamp_id>')
def generate_timestamp_pdf(timestamp_id):
    """Générer un PDF de certification pour un horodatage"""
    if 'username' not in session:
        flash("Veuillez vous connecter pour générer un certificat PDF.", "error")
        return redirect(url_for('login'))
    
    # Récupérer l'horodatage depuis la base de données
    timestamp = Timestamp.query.get_or_404(timestamp_id)
    
    # Vérifier que l'utilisateur a le droit d'accéder à cet horodatage
    if timestamp.user_id and timestamp.user_id != session.get('user_id'):
        flash("Vous n'êtes pas autorisé à accéder à ce certificat.", "error")
        return redirect(url_for('timestamp_service_page'))
    
    # Créer un fichier PDF temporaire
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
        tmp_path = tmp_file.name
    
    # Configurer le document PDF
    doc = SimpleDocTemplate(
        tmp_path, 
        pagesize=A4,
        title=f"Certificat d'horodatage - {timestamp.timestamp_value.strftime('%Y-%m-%d %H:%M:%S')}",
        author="TrustSecure Services",
        subject="Certificat d'horodatage RFC 3161"
    )
    
    # Styles pour le PDF
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=1,
        spaceAfter=12
    )
    section_style = styles["Heading2"]
    normal_style = styles["Normal"]
    
    # Contenu du PDF
    elements = []
    
    # Titre
    elements.append(Paragraph("CERTIFICAT D'HORODATAGE", title_style))
    elements.append(Spacer(1, 20))
    
    
    try:
        logo_path = os.path.join('static', 'img', 'logo.png')
        if os.path.exists(logo_path):
            logo = Image(logo_path, width=150, height=50)
            elements.append(logo)
            elements.append(Spacer(1, 10))
    except:
        pass
    
    
    elements.append(Paragraph("Informations générales", section_style))
    elements.append(Spacer(1, 6))
    
    general_info = [
        ["Date de génération:", datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")],
        ["Référence:", f"TS-{timestamp.id}-{timestamp.timestamp_value.strftime('%Y%m%d%H%M%S')}"],
        ["Émis par:", "TrustSecure Services"]
    ]
    
    t = Table(general_info, colWidths=[150, 300])
    t.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 15))
    
    # Informations sur la donnée horodatée
    elements.append(Paragraph("Données horodatées", section_style))
    elements.append(Spacer(1, 6))
    
    data_info = [
        ["Description:", timestamp.data_description],
        ["Hash (Empreinte):", timestamp.data_hash],
        ["Algorithme de hash:", timestamp.hash_algorithm]
    ]
    
    t = Table(data_info, colWidths=[150, 300])
    t.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 15))
    
    # Informations sur l'horodatage
    elements.append(Paragraph("Horodatage", section_style))
    elements.append(Spacer(1, 6))
    
    timestamp_info = [
        ["Date et heure (UTC):", timestamp.timestamp_value.strftime("%d/%m/%Y %H:%M:%S.%f")[:-3]],
        ["Jeton d'horodatage:", timestamp.tsa_token_sim]
    ]
    
    t = Table(timestamp_info, colWidths=[150, 300])
    t.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 20))
    
    # Explication
    explanation = """Ce certificat d'horodatage atteste de l'existence de la donnée décrite ci-dessus à la date et l'heure indiquées. 
    L'horodatage est conforme à la norme RFC 3161 et a été effectué par un service d'horodatage qualifié.
    L'intégrité de ce certificat peut être vérifiée à l'aide du jeton d'horodatage."""
    
    elements.append(Paragraph(explanation, normal_style))
    elements.append(Spacer(1, 30))
    
    # Avertissement
    disclaimer = """Avertissement : Ce certificat est destiné à des fins informatives uniquement. 
    Pour une preuve légale, veuillez consulter un juriste ou un expert en conformité réglementaire."""
    
    disclaimer_style = ParagraphStyle(
        'Disclaimer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey
    )
    elements.append(Paragraph(disclaimer, disclaimer_style))
    
    # Générer le PDF
    doc.build(elements)
    
    # Renvoyer le fichier PDF généré
    return send_file(
        tmp_path,
        as_attachment=True,
        download_name=f"certificat_horodatage_{timestamp.id}.pdf",
        mimetype='application/pdf'
    )

@app.route('/service/timestamp/download_token/<int:timestamp_id>')
def download_timestamp_token(timestamp_id):
    """Télécharger le jeton d'horodatage brut"""
    if 'username' not in session:
        flash("Veuillez vous connecter pour télécharger ce jeton.", "error")
        return redirect(url_for('login'))
    
    # Récupérer l'horodatage depuis la base de données
    timestamp = Timestamp.query.get_or_404(timestamp_id)
    
    # Vérifier que l'utilisateur a le droit d'accéder à cet horodatage
    if timestamp.user_id and timestamp.user_id != session.get('user_id'):
        flash("Vous n'êtes pas autorisé à accéder à ce jeton.", "error")
        return redirect(url_for('timestamp_service_page'))
    
    # Récupérer le jeton (peut être en base64 ou un format simulé)
    token_content = timestamp.tsa_token_sim
    
    # Renvoyer le jeton comme fichier
    return send_file(
        io.BytesIO(token_content.encode('utf-8')),
        as_attachment=True,
        download_name=f"token_{timestamp.id}.tsr",
        mimetype='application/octet-stream'
    )

@app.route('/service/timestamp/verify', methods=['GET', 'POST'])
def verify_timestamp():
    """Vérifier un jeton d'horodatage existant"""
    result = None
    error = None
    
    if request.method == 'POST':
        try:
            # Récupérer les données de formulaire
            data_hash = request.form.get('data_hash')
            token_data = request.form.get('token_data')
            
            if not data_hash or not token_data:
                error = "Le hash et le jeton d'horodatage sont requis pour la vérification."
                return render_template('timestamp_verify.html', error=error)
            
            # Convertir le hash en bytes
            try:
                binary_hash = bytes.fromhex(data_hash)
            except ValueError:
                error = "Format de hash invalide. Veuillez fournir un hash hexadécimal valide."
                return render_template('timestamp_verify.html', error=error)
            
            # Déterminer si le jeton est en base64 ou brut
            try:
                if token_data.strip().startswith("TSA_TOKEN_SIM_"):
                    # Token simulé ancien format
                    result = {
                        "valid": True,
                        "simulated": True,
                        "timestamp": "Simulé (non vérifié)",
                        "message": "Ce jeton est un jeton de simulation et ne fournit pas une preuve d'horodatage avec la même force qu'un jeton RFC 3161 officiel."
                    }
                else:
                    # Essayer de décoder le jeton comme base64
                    try:
                        # Décoder le jeton base64
                        token_bytes = base64.b64decode(token_data)
                        
                        # Essayer de parser en tant que JSON (format simulation)
                        try:
                            token_json = json.loads(token_bytes.decode('utf-8'))
                            
                            # Vérifier si c'est notre format de jeton structuré
                            if 'version' in token_json and 'hash_value' in token_json and 'tsa_signature' in token_json:
                                # C'est notre format de jeton structuré
                                token_hash = token_json['hash_value']
                                token_time = token_json['timestamp']
                                token_sig = base64.b64decode(token_json['tsa_signature'])
                                
                                # Vérifier que le hash correspond
                                if token_hash != data_hash:
                                    result = {
                                        "valid": False,
                                        "message": "Le hash dans le jeton ne correspond pas au hash fourni",
                                        "details": f"Hash attendu: {data_hash}\nHash dans le jeton: {token_hash}"
                                    }
                                else:
                                    # Vérifier la signature
                                    try:
                                        # Reconstruire les données signées
                                        timestamp_data = f"{data_hash}:{token_time}"
                                        
                                        # Vérifier la signature avec la clé publique
                                        public_key.verify(
                                            token_sig,
                                            timestamp_data.encode(),
                                            padding.PSS(
                                                mgf=padding.MGF1(hashes.SHA256()),
                                                salt_length=padding.PSS.MAX_LENGTH
                                            ),
                                            hashes.SHA256()
                                        )
                                        
                                        # Si nous arrivons ici, la vérification a réussi
                                        parsed_time = parser.parse(token_time)
                                        result = {
                                            "valid": True,
                                            "simulated": True,
                                            "timestamp": parsed_time.strftime("%Y-%m-%d %H:%M:%S.%f")[:-3] + "Z",
                                            "certificate": "Vérifié par TrustSecure Demo TSA",
                                            "message": "Le jeton a été vérifié avec succès. Le document existait au moment indiqué."
                                        }
                                    except Exception as sig_error:
                                        result = {
                                            "valid": False,
                                            "message": "La signature du jeton est invalide",
                                            "details": str(sig_error)
                                        }
                            else:
                                # Ce n'est pas notre format JSON, message d'erreur spécifique
                                result = {
                                    "valid": False,
                                    "message": "Format de jeton non reconnu",
                                    "details": "Le jeton n'est pas au format TrustSecure."
                                }
                        except json.JSONDecodeError:
                            # Ce n'est pas du JSON, message d'erreur spécifique
                            result = {
                                "valid": False,
                                "message": "Format de jeton non reconnu",
                                "details": "Le jeton n'est pas au format JSON attendu."
                            }
                    except Exception as b64_error:
                        # Problème avec le décodage base64
                        result = {
                            "valid": False,
                            "message": "Le jeton n'est pas au format base64 valide",
                            "details": str(b64_error)
                        }
            
            except Exception as e:
                error = f"Erreur lors de la vérification du jeton: {str(e)}"
        
        except Exception as general_error:
            error = f"Erreur générale: {str(general_error)}"
    
    return render_template('timestamp_verify.html', result=result, error=error)

@app.route('/download_public_key')
def download_public_key_pem():
    return send_file(io.BytesIO(pem_public_key), as_attachment=True, download_name='public_key.pem', mimetype='text/plain')

if __name__ == '__main__':
    # Create the database tables if they don't exist
    with app.app_context():
        # Drop all tables first to ensure a clean state
        db.drop_all()
        # Create all tables with current models
        db.create_all()
        print("Database tables dropped and recreated successfully.")
    app.run(debug=True)